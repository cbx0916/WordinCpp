from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Inches
import sys

# 使用这个函数生成文本
def addText(titlename, text, dest):
    doc = Document() #新的doc文本

    doc.core_properties.author="author" # 设置author
    
    # title = doc.add_heading("红头文件",level=0) # 设置标题
    title = doc.add_paragraph("XXXXXX公司") # 设置标题
    run_title = title.runs[0]
    run_title.bold = True # 加粗
    run_title.font.size = Pt(26) # 字号
    run_title.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
    run_title.font.name = "黑体"
    run_title.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中

    outgoing = doc.add_paragraph("xx字【20xx】x号") # 设置发文字号
    run_out = outgoing.runs[0]
    run_out.bold = True # 加粗
    run_out.font.size = Pt(14) # 字号
    run_out.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_out.font.name = "仿宋"
    run_out.element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    outgoing.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中

    # red_line = doc.add_picture('redline.png')
    red_line = doc.add_paragraph()
    red_line.paragraph_format.line_spacing = 1.5  # 行距1.5倍
    red_line.paragraph_format.first_line_indent = 0 # 缩进为0
    red_line.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中
    red_line.add_run().add_picture('redline.png',width=Inches(6.2)) # 长度15.5cm


    stitle = doc.add_paragraph(titlename) # 设置标题
    run_stitle = stitle.runs[0]
    run_stitle.bold = True # 加粗
    run_stitle.font.size = Pt(16) # 字号
    run_stitle.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run_stitle.font.name = "黑体"
    run_stitle.element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    stitle.alignment = WD_ALIGN_PARAGRAPH.CENTER # 居中

    paragraph = doc.add_paragraph(text) # 设置正文
    run_para = paragraph.runs[0]
    run_para.font.name = "仿宋"
    run_para.element.rPr.rFonts.set(qn('w:eastAsia'),'仿宋')
    run_para.font.size = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.5  # 行距1.5倍，固定值28磅为Pt(28)
    paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
    paragraph.paragraph_format.first_line_indent = Pt(28) # 缩进两个字符

    doc.save(dest)
    

# addText("关于下发《XXXX》的通知", "正文......", "dest.docx")

# 删除段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None
 

def setMargin(docx):
    section = docx.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)


# 使用这个函数修改文本
def dealParagraph(file, para_th, text, newtext):
    doc = Document(file) #读取路径下的docx
    print(len(doc.paragraphs)) # 获取doSc段落数量，标题属于段落
    if para > len(doc.paragraphs):
        print("out length")
        return
    doc.paragraphs[para_th].text = text # 修改第para_th段的文本
    doc.add_paragraph(newtext) # 添加新正文
    doc.save("test1.docx")


# doc.add_paragraph("text") 添加新段落
# paragraph.text 获取段落的文本
# paragraph.style = "stylename" 设置段落样式
# doc.paragraphs 获取全部段落
# doc.paragraphs[0].text 第一段的文本





# def create_docx(text):
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save('output.docx')

# if __name__ == "__main__":
#     print("")
#     if len(sys.argv) < 2:
#         print("Usage: python create_docx.py 'your text here'")
#         sys.exit(1)
    
#     text = sys.argv[1]
#     print(text)
#     create_docx(text)